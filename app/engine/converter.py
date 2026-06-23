"""
Conversion pipeline: PDF → DOCX or MD.
Runs in a worker thread; communicates progress via Qt signals.
"""
import os
import re
import shutil
import tempfile
from pathlib import Path
from typing import Callable, Optional

from PyQt6.QtCore import QObject, pyqtSignal, QRunnable, QThreadPool

import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt


def _slugify(name: str) -> str:
    """Convert filename stem to lowercase-hyphenated form."""
    stem = Path(name).stem
    slug = stem.lower()
    slug = re.sub(r"[\s_]+", "-", slug)
    slug = re.sub(r"[^a-z0-9\-]", "", slug)
    slug = re.sub(r"-{2,}", "-", slug).strip("-")
    return slug


def build_output_path(source_path: str, output_dir: Optional[str], fmt: str) -> str:
    src = Path(source_path)
    slug = _slugify(src.name)
    filename = f"{slug}-converted.{fmt}"
    base = Path(output_dir) if output_dir else src.parent
    return str(base / filename)


def has_text_layer(pdf_path: str) -> bool:
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:3]:
                text = page.extract_text() or ""
                if len(text.strip()) > 20:
                    return True
    except Exception:
        pass
    return False


def extract_text_from_pdf(pdf_path: str, poppler_path: Optional[str] = None,
                           progress_cb: Optional[Callable[[float], None]] = None) -> list[str]:
    """Return list of page text strings."""
    if has_text_layer(pdf_path):
        pages_text = []
        with pdfplumber.open(pdf_path) as pdf:
            total = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                pages_text.append(page.extract_text() or "")
                if progress_cb:
                    progress_cb((i + 1) / total * 0.8)
        return pages_text
    else:
        # Scanned: rasterise + OCR
        kwargs = {}
        if poppler_path:
            kwargs["poppler_path"] = poppler_path
        images = convert_from_path(pdf_path, dpi=300, **kwargs)
        total = len(images)
        pages_text = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang="eng")
            pages_text.append(text)
            if progress_cb:
                progress_cb((i + 1) / total * 0.8)
        return pages_text


def _structure_paragraphs(pages_text: list[str]) -> list[dict]:
    """
    Very lightweight structuring: detect headings by ALL-CAPS short lines
    or lines ending without punctuation that are short. Everything else = paragraph.
    Returns list of {type: heading|para, text: str}.
    """
    blocks = []
    for page_text in pages_text:
        if not page_text.strip():
            continue
        lines = page_text.splitlines()
        buffer = []
        for raw in lines:
            line = raw.strip()
            if not line:
                if buffer:
                    blocks.append({"type": "para", "text": " ".join(buffer)})
                    buffer = []
                continue
            is_heading = (
                len(line) < 80
                and (line.isupper() or (line == line.title() and not line[-1] in ".,:;?!"))
                and not line.startswith("-")
                and len(line.split()) <= 10
            )
            if is_heading:
                if buffer:
                    blocks.append({"type": "para", "text": " ".join(buffer)})
                    buffer = []
                blocks.append({"type": "heading", "text": line})
            else:
                buffer.append(line)
        if buffer:
            blocks.append({"type": "para", "text": " ".join(buffer)})
    return blocks


def write_docx(pages_text: list[str], output_path: str) -> None:
    doc = Document()
    blocks = _structure_paragraphs(pages_text)
    for block in blocks:
        if block["type"] == "heading":
            h = doc.add_heading(block["text"], level=1)
            h.style.font.size = Pt(14)
        else:
            doc.add_paragraph(block["text"])
    doc.save(output_path)


def write_md(pages_text: list[str], output_path: str) -> None:
    blocks = _structure_paragraphs(pages_text)
    lines = []
    for block in blocks:
        if block["type"] == "heading":
            lines.append(f"## {block['text']}\n")
        else:
            lines.append(f"{block['text']}\n")
        lines.append("")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


class ConvertSignals(QObject):
    file_progress = pyqtSignal(int, float)       # index, 0–1
    file_done = pyqtSignal(int, str)             # index, output_path
    file_error = pyqtSignal(int, str)            # index, error_message
    overall_progress = pyqtSignal(float)         # 0–1
    batch_complete = pyqtSignal()
    log_line = pyqtSignal(str, str)              # level, message


class ConvertWorker(QRunnable):
    def __init__(self, files: list, fmt: str, output_dir: Optional[str],
                 poppler_path: Optional[str] = None):
        super().__init__()
        self.files = files          # list of file paths
        self.fmt = fmt
        self.output_dir = output_dir
        self.poppler_path = poppler_path
        self.signals = ConvertSignals()
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def run(self):
        total = len(self.files)
        for idx, src_path in enumerate(self.files):
            if self._cancelled:
                break

            self.signals.log_line.emit("info", f"Starting: {Path(src_path).name}")
            output_path = build_output_path(src_path, self.output_dir, self.fmt)

            def file_progress(p: float):
                self.signals.file_progress.emit(idx, p)
                overall = (idx + p) / total
                self.signals.overall_progress.emit(overall)

            try:
                pages_text = extract_text_from_pdf(
                    src_path,
                    poppler_path=self.poppler_path,
                    progress_cb=file_progress,
                )
                if not any(t.strip() for t in pages_text):
                    self.signals.log_line.emit(
                        "error",
                        f"WARN: {Path(src_path).name} produced empty text — check the output file.",
                    )

                if self.fmt == "docx":
                    write_docx(pages_text, output_path)
                else:
                    write_md(pages_text, output_path)

                self.signals.file_progress.emit(idx, 1.0)
                self.signals.file_done.emit(idx, output_path)
                self.signals.log_line.emit("success", f"Done: {Path(output_path).name}")
            except Exception as e:
                msg = _human_error(str(e))
                self.signals.file_error.emit(idx, msg)
                self.signals.log_line.emit("error", f"SKIPPED: {Path(src_path).name} — {msg}")

            overall = (idx + 1) / total
            self.signals.overall_progress.emit(overall)

        self.signals.batch_complete.emit()


def _human_error(raw: str) -> str:
    raw_lower = raw.lower()
    if "password" in raw_lower or "encrypted" in raw_lower:
        return "Could not read this file — it may be password-protected."
    if "permission" in raw_lower or "access" in raw_lower:
        return "Could not write the output — check folder permissions."
    if "tesseract" in raw_lower:
        return "Tesseract OCR is not installed or not on PATH."
    if "poppler" in raw_lower or "pdftoppm" in raw_lower:
        return "Poppler utilities are missing — cannot rasterise scanned pages."
    if "no such file" in raw_lower or "not found" in raw_lower:
        return "Could not read this file — it may have been moved or deleted."
    if raw:
        return f"An unexpected error occurred: {raw[:120]}"
    return "An unexpected error occurred."
